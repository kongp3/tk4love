# -*- coding: utf-8 -*-

import docx


if __name__ == '__main__':

    rp_list = [
        {'name': u'孔繁1', 'mobile': '18610812902', 'email': 'kongp6@outlook.com'},
        {'name': u'孔繁2', 'mobile': '18610812903', 'email': 'kongp7@outlook.com'},
        {'name': u'孔繁3', 'mobile': '18610812904', 'email': 'kongp8@outlook.com'},
        {'name': u'孔繁4', 'mobile': '18610812905', 'email': 'kongp9@outlook.com'},
    ]

    name = u'标题'
    root = '/Users/kongfm/MyPMP/tk4love/doc/'
    path = root + 'template.docx'

    for rp in rp_list:
        # 1. 初始化一个文档
        doc = docx.Document(path)

        # 2. 遍历并处理文档中的段落
        for para in doc.paragraphs:

            # 2.1 处理文本的替换(除去有非文本内容的段落)
            if '$NAME$' in para.text or '$MOBILE$' in para.text:
                run = para.runs[0]
                para.text = para.text.replace('$NAME$', rp.get('name'))
                para.text = para.text.replace('$MOBILE$', rp.get('mobile'))
                new_run = para.runs[0]
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size

            # 2.2 处理有非文本内容的段落(直接赋值会把非文本弄丢)
            # 非文本--红头文件的红头、公章、图片等等
            if '$EMAIL$' in para.text:

                # 2.2.1 获取所有文本段儿
                runs = para.runs

                # 2.2.2 准备temp文本段儿
                run_list = []

                # 2.2.3 保存初始文本段儿信息
                run_tmp = para.runs[0]

                # 2.2.4 遍历处理文本段儿
                for run in para.runs:

                    # 2.2.4.1 如果有文本做替换, 放入temp, 清空原有文本的操作
                    if run.text:
                        run.text = run.text.replace('$EMAIL$', rp.get('email'))
                        run_list.append(run.text)
                        run.text = ''
                # 2.2.5 将处理过的文本段儿放入段落
                para.add_run(text=''.join(run_list))

                # 2.2.6 获取段落中新的文本段儿
                new_run = para.runs[-1]

                # 设置字体、大小等信息(要调试获取才可得)
                new_run.font.name = run_tmp.font.name
                new_run.font.size = run_tmp.font.size

                # # 获取文字格式信息
                # print u'字体名称：', run.font.name
                # print u'字体名称：', new_run.font.name
                # # 字体名称： 宋体
                # print u'字体大小：', run.font.size
                # print u'字体大小：', new_run.font.size
                # # 字体大小： 152400
                # # print u'是否加粗：', run.font.bold
                # # # 是否加粗： None
                # # print u'是否斜体：', run.font.italic
                # # # 是否斜体： True
                # # print u'字体颜色：', run.font.color.rgb
                # # # 字体颜色： FF0000
                # # print u'字体高亮：', run.font.highlight_color
                # # # 字体高亮： YELLOW (7)
                # # print u'下划线：', run.font.underline
                # # # 下划线： True
                # # print u'删除线：', run.font.strike
                # # # 删除线： None
                # # print u'双删除线：', run.font.double_strike
                # # # 双删除线： None
                # # print u'下标：', run.font.subscript
                # # # 下标： None
                # # print u'上标：', run.font.superscript

        # 3. 保存处理后的文档
        doc.save(root + name + '-' + rp.get('name') + '.docx')
